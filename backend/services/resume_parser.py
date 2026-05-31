"""
Resume Parser Service.

Extracts raw text from uploaded resume files (PDF, DOCX).
"""
import io
import logging

from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


class ResumeParser:
    """Extracts text content from resume files."""

    @staticmethod
    def parse(file_obj, content_type: str) -> str:
        """
        Parse a resume file and return extracted text.

        Args:
            file_obj: Django UploadedFile or file-like object.
            content_type: MIME type of the file.

        Returns:
            Extracted text as a string.
        """
        if content_type == 'application/pdf':
            return ResumeParser._parse_pdf(file_obj)
        elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return ResumeParser._parse_docx(file_obj)
        else:
            raise ValueError(f"Unsupported file type: {content_type}")

    @staticmethod
    def _parse_pdf(file_obj) -> str:
        """Extract text from a PDF file."""
        try:
            reader = PdfReader(io.BytesIO(file_obj.read()))
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            text = '\n'.join(text_parts)
            logger.info(f"Extracted {len(text)} characters from PDF ({len(reader.pages)} pages)")
            return text.strip()
        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            raise ValueError(f"Failed to parse PDF: {e}")

    @staticmethod
    def _parse_docx(file_obj) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = Document(io.BytesIO(file_obj.read()))
            text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
            text = '\n'.join(text_parts)
            logger.info(f"Extracted {len(text)} characters from DOCX ({len(doc.paragraphs)} paragraphs)")
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            raise ValueError(f"Failed to parse DOCX: {e}")
